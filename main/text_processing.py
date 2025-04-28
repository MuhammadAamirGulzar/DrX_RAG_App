import os
from typing import List, Dict, Any, Tuple, Union, Optional
import torch
import tiktoken
import tempfile
import PyPDF2
import docx
import pandas as pd
import time

# Import the Excel/CSV processing code
from main.excel_processing import (
    read_excel, read_csv, process_tabular_file, 
    RowBasedChunking, ColumnBasedChunking, SemanticChunking,
    infer_best_chunking_strategy
)

# File handling functions
def read_pdf(file_path: str) -> List[Tuple[int, str]]:
    """
    Read PDF file and extract text with page numbers.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        List of tuples (page_number, text)
    """
    pages = []
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for i, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text:
                pages.append((i+1, text))
    return pages

def read_docx(file_path: str) -> List[Tuple[int, str]]:
    """
    Read DOCX file and extract text with page numbers.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of tuples (page_number, text)
    """
    doc = docx.Document(file_path)
    # DOCX doesn't have native page numbers, so we'll use section breaks as a proxy
    full_text = ""
    pages = []
    page_num = 1
    
    for para in doc.paragraphs:
        if para.text.strip() == "" and full_text:  # Consider empty paragraphs as potential page breaks
            pages.append((page_num, full_text))
            full_text = ""
            page_num += 1
        else:
            full_text += para.text + "\n"
    
    # Add any remaining text
    if full_text:
        pages.append((page_num, full_text))
    
    # If no page breaks were detected, just return everything as page 1
    if not pages:
        full_text = "\n".join([para.text for para in doc.paragraphs])
        pages.append((1, full_text))
    
    return pages

def read_file(file_path: str) -> List[Tuple[int, str]]:
    """
    Read file based on extension and extract text with page numbers.
    
    Args:
        file_path: Path to the file
        
    Returns:
        List of tuples (page_number, text)
    """
    if file_path.lower().endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return read_docx(file_path)
    elif file_path.lower().endswith(('.xlsx', '.xls')):
        return read_excel(file_path)
    elif file_path.lower().endswith(('.csv')):
        return read_csv(file_path)
        # return [(0, "Tabular file - please use process_tabular_file instead")]
    else:
        raise ValueError(f"Unsupported file format: {file_path}")

def chunk_text(text: str, filename: str, page_num: int, max_tokens: int = 512) -> List[Dict]:
    """
    Split text into token-limited chunks.
    
    Args:
        text: Text to split
        filename: Name of the source file
        page_num: Page number of the source
        max_tokens: Maximum tokens per chunk
        
    Returns:
        List of chunk dictionaries
    """
    tokenizer = tiktoken.get_encoding("cl100k_base")
    tokens = tokenizer.encode(text)
    chunks = []
    
    for i in range(0, len(tokens), max_tokens):
        chunk_tokens = tokens[i:i+max_tokens]
        chunk_text = tokenizer.decode(chunk_tokens)
        chunks.append({
            "chunk_number": len(chunks) + 1,
            "text": chunk_text,
            "meta_data": {
                "filename": filename, 
                "page_number": page_num,
                "content_type": "document_text"
            }
        })
    
    return chunks

def process_file(file_path: str, filename: str, 
                max_tokens: int = 512, 
                chunking_strategy: str = "auto") -> List[Dict]:
    """
    Process a file and create chunks.
    
    Args:
        file_path: Path to the file
        filename: Name to use in metadata
        max_tokens: Maximum tokens per chunk
        chunking_strategy: Strategy for tabular data ("row", "column", "semantic", or "auto")
        
    Returns:
        List of chunks
    """
    # Process based on file extension
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext in ['.xlsx', '.xls', '.csv']:
        # For tabular data files
        
        # Select chunking strategy
        strategy_class = RowBasedChunking  # Default
        group_by_column = None
        
        if chunking_strategy == "auto":
            # Infer the best strategy based on data characteristics
            strategy_class, group_by_column = infer_best_chunking_strategy(file_path)
        elif chunking_strategy == "row":
            strategy_class = RowBasedChunking
        elif chunking_strategy == "column":
            strategy_class = ColumnBasedChunking
        elif chunking_strategy == "semantic":
            strategy_class, group_by_column = infer_best_chunking_strategy(file_path)
            strategy_class = SemanticChunking
            
        # Process the tabular file
        chunks, structure_summary = process_tabular_file(
            file_path, 
            filename, 
            chunking_strategy=strategy_class,
            group_by_column=group_by_column,
            max_tokens=max_tokens
        )
        
        # Add the structure summary to the chunks
        all_chunks = [structure_summary] + chunks
        
    else:
        # Standard processing for PDF and DOCX
        all_chunks = []
        pages = read_file(file_path)
        
        for page_num, text in pages:
            all_chunks.extend(chunk_text(text, filename, page_num, max_tokens))
    
    return all_chunks

def get_file_type(file_name: str) -> str:
    """
    Determine file type from file name
    
    Args:
        file_name: Name of the file
        
    Returns:
        File type description
    """
    extension = os.path.splitext(file_name)[1].lower()
    
    if extension == '.pdf':
        return "PDF Document"
    elif extension == '.docx':
        return "Word Document"
    elif extension in ['.xlsx', '.xls']:
        return "Excel Spreadsheet"
    elif extension == '.csv':
        return "CSV File"
    else:
        return "Unknown File Type"
