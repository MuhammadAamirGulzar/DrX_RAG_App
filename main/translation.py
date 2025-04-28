import os
from typing import List, Dict, Any, Tuple, Union, Optional
import torch
import tempfile
import streamlit as st
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer, MarianMTModel, MarianTokenizer
import docx
import PyPDF2
import pandas as pd
import time

# Map of language codes to full names
LANGUAGE_MAP = {
    "ar": "Arabic",
    "de": "German",
    "es": "Spanish",
    "fr": "French",
    "hi": "Hindi",
    "it": "Italian",
    "ja": "Japanese",
    "ko": "Korean",
    "pt": "Portuguese",
    "ru": "Russian",
    "tr": "Turkish",
    "ur": "Urdu",
    "zh": "Chinese",
    "en": "English"
}

# Reverse mapping
LANGUAGE_CODE_MAP = {v: k for k, v in LANGUAGE_MAP.items()}

def get_translation_model(source_lang: str, target_lang: str = "en"):
    """
    Load the appropriate translation model based on source and target languages.
    
    Args:
        source_lang: Source language code (e.g., 'fr' for French)
        target_lang: Target language code (default: 'en' for English)
        
    Returns:
        tuple: (model, tokenizer)
    """
    # Check if we're already trying to translate to English
    if source_lang == "en":
        # No translation needed
        return None, None
    
    # Use Helsinki-NLP/opus-mt models for translation
    if target_lang == "en":
        model_name = f"Helsinki-NLP/opus-mt-{source_lang}-{target_lang}"
    else:
        model_name = f"Helsinki-NLP/opus-mt-{source_lang}-{target_lang}"
    
    try:
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)
        return model, tokenizer
    except Exception as e:
        # Fallback to a multilingual model if specific pair isn't available
        st.warning(f"Specific translation model not found: {e}. Using multilingual fallback.")
        model_name = "facebook/mbart-large-50-many-to-many-mmt"
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        return model, tokenizer

def translate_text(text: str, model, tokenizer, source_lang: str, target_lang: str = "en", max_length: int = 512) -> str:
    """
    Translate a piece of text using the provided model.
    
    Args:
        text: Text to translate
        model: Translation model
        tokenizer: Tokenizer for the model
        source_lang: Source language code
        target_lang: Target language code (default: English)
        max_length: Maximum length of translated text
        
    Returns:
        str: Translated text
    """
    # No translation needed if source is already the target language
    if source_lang == target_lang:
        return text
    
    # Handle empty text
    if not text.strip():
        return text
    
    # Split long text into paragraphs to handle within model's context window
    paragraphs = text.split('\n')
    translated_paragraphs = []
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            translated_paragraphs.append("")
            continue
            
        # Encode the text
        encoded = tokenizer(paragraph, return_tensors="pt", padding=True, truncation=True, max_length=max_length)
        
        # Set the language tags for some models that require it
        if hasattr(tokenizer, "set_src_lang_special_tokens"):
            tokenizer.set_src_lang_special_tokens(source_lang)
        
        # Generate translation
        with torch.no_grad():
            output = model.generate(**encoded, max_length=max_length, num_beams=4, early_stopping=True)
        
        # Decode the output
        translated_text = tokenizer.decode(output[0], skip_special_tokens=True)
        translated_paragraphs.append(translated_text)
    
    # Join paragraphs back with newlines to preserve document structure
    return '\n'.join(translated_paragraphs)

def translate_pdf(file_path: str, translated_file_path: str, model, tokenizer, source_lang: str, target_lang: str = "en") -> bool:
    """
    Translate a PDF file and save the translated version.
    
    Args:
        file_path: Path to the source PDF file
        translated_file_path: Path to save the translated PDF
        model: Translation model
        tokenizer: Tokenizer for the model
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        bool: Success status
    """
    try:
        # Read the PDF
        pages = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        
        # Translate each page
        translated_pages = []
        for i, page_text in enumerate(pages):
            with st.spinner(f"Translating page {i+1}/{len(pages)}..."):
                translated_text = translate_text(page_text, model, tokenizer, source_lang, target_lang)
                translated_pages.append(translated_text)
        
        # Create a new PDF with translated text
        from fpdf import FPDF
        pdf = FPDF()
        
        for page_text in translated_pages:
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Split text by lines and add to PDF
            lines = page_text.split('\n')
            for line in lines:
                pdf.multi_cell(190, 10, txt=line)
        
        # Save the translated PDF
        pdf.output(translated_file_path)
        return True
        
    except Exception as e:
        st.error(f"Error translating PDF: {str(e)}")
        return False

def translate_docx(file_path: str, translated_file_path: str, model, tokenizer, source_lang: str, target_lang: str = "en") -> bool:
    """
    Translate a DOCX file and save the translated version.
    
    Args:
        file_path: Path to the source DOCX file
        translated_file_path: Path to save the translated DOCX
        model: Translation model
        tokenizer: Tokenizer for the model
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        bool: Success status
    """
    try:
        # Read the DOCX
        doc = docx.Document(file_path)
        
        # Create a new document for the translation
        translated_doc = docx.Document()
        
        # Translate each paragraph and add to the new document
        total_paragraphs = len(doc.paragraphs)
        
        for i, para in enumerate(doc.paragraphs):
            if i % 5 == 0:  # Update progress every 5 paragraphs
                st.spinner(f"Translating paragraph {i+1}/{total_paragraphs}...")
                
            if para.text.strip():  # Skip empty paragraphs
                translated_text = translate_text(para.text, model, tokenizer, source_lang, target_lang)
                translated_para = translated_doc.add_paragraph(translated_text)
                
                # Try to preserve basic formatting
                if para.style:
                    translated_para.style = para.style
            else:
                translated_doc.add_paragraph()  # Add empty paragraph to preserve structure
        
        # Save the translated document
        translated_doc.save(translated_file_path)
        return True
        
    except Exception as e:
        st.error(f"Error translating DOCX: {str(e)}")
        return False

def translate_excel(file_path: str, translated_file_path: str, model, tokenizer, source_lang: str, target_lang: str = "en") -> bool:
    """
    Translate an Excel file and save the translated version.
    
    Args:
        file_path: Path to the source Excel file
        translated_file_path: Path to save the translated Excel
        model: Translation model
        tokenizer: Tokenizer for the model
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        bool: Success status
    """
    try:
        # Get all sheets from the Excel file
        excel_file = pd.ExcelFile(file_path)
        
        # Create a new Excel writer
        with pd.ExcelWriter(translated_file_path, engine='openpyxl') as writer:
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                # Read the sheet
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Translate column names
                translated_columns = []
                for col in df.columns:
                    translated_col = translate_text(str(col), model, tokenizer, source_lang, target_lang)
                    translated_columns.append(translated_col)
                
                # Create a new DataFrame with translated column names
                translated_df = pd.DataFrame(columns=translated_columns)
                
                # Translate cell values (only string/object columns)
                total_rows = len(df)
                for i, row in enumerate(df.itertuples(index=False)):
                    if i % 10 == 0:  # Update progress every 10 rows
                        st.spinner(f"Translating sheet '{sheet_name}', row {i+1}/{total_rows}...")
                    
                    row_data = []
                    for j, val in enumerate(row):
                        if pd.isna(val):
                            row_data.append(val)
                        elif df.dtypes[j] == 'object':
                            # Only translate text data
                            translated_val = translate_text(str(val), model, tokenizer, source_lang, target_lang)
                            row_data.append(translated_val)
                        else:
                            # Keep non-text data as is
                            row_data.append(val)
                    
                    translated_df.loc[i] = row_data
                
                # Write the translated DataFrame to the output file
                translated_df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        return True
        
    except Exception as e:
        st.error(f"Error translating Excel: {str(e)}")
        return False

def translate_csv(file_path: str, translated_file_path: str, model, tokenizer, source_lang: str, target_lang: str = "en") -> bool:
    """
    Translate a CSV file and save the translated version.
    
    Args:
        file_path: Path to the source CSV file
        translated_file_path: Path to save the translated CSV
        model: Translation model
        tokenizer: Tokenizer for the model
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        bool: Success status
    """
    try:
        # Read the CSV file
        try:
            df = pd.read_csv(file_path)
        except UnicodeDecodeError:
            # Try different encodings if UTF-8 fails
            df = pd.read_csv(file_path, encoding='latin1')
        except pd.errors.ParserError:
            # Try different delimiters if comma fails
            try:
                df = pd.read_csv(file_path, sep=';')
            except:
                df = pd.read_csv(file_path, sep='\t')
        
        # Translate column names
        translated_columns = []
        for col in df.columns:
            translated_col = translate_text(str(col), model, tokenizer, source_lang, target_lang)
            translated_columns.append(translated_col)
        
        # Create a new DataFrame with translated column names
        translated_df = pd.DataFrame(columns=translated_columns)
        
        # Translate cell values (only string/object columns)
        total_rows = len(df)
        for i, row in enumerate(df.itertuples(index=False)):
            if i % 10 == 0:  # Update progress every 10 rows
                st.spinner(f"Translating row {i+1}/{total_rows}...")
            
            row_data = []
            for j, val in enumerate(row):
                if pd.isna(val):
                    row_data.append(val)
                elif df.dtypes[j] == 'object':
                    # Only translate text data
                    translated_val = translate_text(str(val), model, tokenizer, source_lang, target_lang)
                    row_data.append(translated_val)
                else:
                    # Keep non-text data as is
                    row_data.append(val)
            
            translated_df.loc[i] = row_data
        
        # Write the translated DataFrame to the output file
        translated_df.to_csv(translated_file_path, index=False)
        return True
        
    except Exception as e:
        st.error(f"Error translating CSV: {str(e)}")
        return False

def translate_document(file_path: str, output_path: str, source_lang: str, target_lang: str = "en") -> Tuple[bool, str]:
    """
    Translate a document from source language to target language.
    
    Args:
        file_path: Path to the source document
        output_path: Path to save the translated document
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        Tuple of (success status, error message if any)
    """
    # Skip translation if source and target are the same
    if source_lang == target_lang:
        # Just copy the file
        try:
            with open(file_path, 'rb') as src, open(output_path, 'wb') as dst:
                dst.write(src.read())
            return True, ""
        except Exception as e:
            return False, str(e)
    
    try:
        # Load translation model
        model, tokenizer = get_translation_model(source_lang, target_lang)
        if model is None:
            # Source is already in target language, just copy
            with open(file_path, 'rb') as src, open(output_path, 'wb') as dst:
                dst.write(src.read())
            return True, ""
        
        # Determine file type and translate accordingly
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            success = translate_pdf(file_path, output_path, model, tokenizer, source_lang, target_lang)
        elif file_ext == '.docx':
            success = translate_docx(file_path, output_path, model, tokenizer, source_lang, target_lang)
        elif file_ext in ['.xlsx', '.xls']:
            success = translate_excel(file_path, output_path, model, tokenizer, source_lang, target_lang)
        elif file_ext == '.csv':
            success = translate_csv(file_path, output_path, model, tokenizer, source_lang, target_lang)
        else:
            return False, f"Unsupported file format: {file_ext}"
        
        if success:
            return True, ""
        else:
            return False, "Translation failed"
            
    except Exception as e:
        return False, str(e)